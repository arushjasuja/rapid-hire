"""Resume text extraction.

Deterministic and LLM-free on purpose. This runs just before the crew so the
intake agent receives plain text rather than a file blob -- piping base64 bytes
through an LLM tool-call would be pointless. Supports the three formats real
resumes actually arrive in: PDF, DOCX, and plain text/markdown.

The single ``extract_text`` entry point accepts whatever the caller has on hand:
a path, raw bytes plus a filename, or a Streamlit ``UploadedFile`` (which exposes
``.name`` and ``.getvalue()``).
"""

from __future__ import annotations

import io
from pathlib import Path

from ..logging import get_logger

log = get_logger(__name__)

_SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown"}


class UnsupportedResumeError(ValueError):
    pass


def extract_text(source, *, filename: str | None = None) -> str:
    """Return the text of a resume from a path, bytes, or an uploaded-file object.

    ``filename`` is only needed when ``source`` is raw bytes (it's how we pick the
    parser). Paths and file-like objects carry their own name.
    """
    data, name = _normalize(source, filename)
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        text = _from_pdf(data)
    elif suffix == ".docx":
        text = _from_docx(data)
    elif suffix in {".txt", ".md", ".markdown"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedResumeError(
            f"Can't parse {name!r}; supported types are {sorted(_SUPPORTED)}"
        )
    # A resume that parsed to almost nothing is usually a scanned/image PDF. We
    # don't OCR here; better to say so than to silently score an empty profile.
    cleaned = text.strip()
    if len(cleaned) < 30:
        log.warning("parsed %r but got only %d chars (image-only PDF?)", name, len(cleaned))
    return cleaned


def _normalize(source, filename: str | None) -> tuple[bytes, str]:
    # Path or path-like string
    if isinstance(source, (str, Path)):
        p = Path(source)
        return p.read_bytes(), p.name
    # Streamlit UploadedFile and similar: has a name and getvalue()
    if hasattr(source, "getvalue") and hasattr(source, "name"):
        return source.getvalue(), source.name
    # Bare file object with read()
    if hasattr(source, "read"):
        name = getattr(source, "name", None) or filename
        if not name:
            raise ValueError("file-like source needs a filename to pick a parser")
        return source.read(), name
    # Raw bytes
    if isinstance(source, (bytes, bytearray)):
        if not filename:
            raise ValueError("bytes source requires filename=")
        return bytes(source), filename
    raise TypeError(f"don't know how to read {type(source).__name__}")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    # extract_text() returns None for pages with no text layer; join what we get.
    pages = (page.extract_text() or "" for page in reader.pages)
    return "\n".join(pages)


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables hold real content on plenty of resumes (skills grids, etc.), so pull
    # cell text too rather than dropping it.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
